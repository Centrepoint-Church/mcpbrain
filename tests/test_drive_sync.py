"""Tests for mcpbrain.sync.drive — fake service, no network."""

import pytest

from mcpbrain.store import Store
from mcpbrain.sync.drive import sync_drive, backfill_drive, normalise_drive, _fetch_text


# ---------------------------------------------------------------------------
# Fake Drive service
# ---------------------------------------------------------------------------

class _Req:
    def __init__(self, result=None, raise_exc=None):
        self._r = result
        self._e = raise_exc

    def execute(self):
        if self._e:
            raise self._e
        return self._r


class _Changes:
    def __init__(self, start_token="100", pages=None, initial_cursor=None):
        self._start = start_token
        self._pages = pages or []
        # The first delta call uses the stored cursor as pageToken.
        # We always route that to pages[0]. Subsequent nextPageToken values
        # are string integers ("1", "2", …) that index directly into _pages.
        self._initial_cursor = initial_cursor  # set by FakeDriveService

    def getStartPageToken(self, **_kw):          # accept driveId/supportsAllDrives
        return _Req({"startPageToken": self._start})

    def list(self, **kw):
        token = kw.get("pageToken")
        if token is None or token == self._initial_cursor:
            # First call in a delta run
            idx = 0
        else:
            try:
                idx = int(token)
            except (ValueError, TypeError):
                idx = 0
        return _Req(self._pages[idx])


class _Files:
    def __init__(self, exports=None, media=None, export_raises=None, file_list=None):
        self._exports = exports or {}
        self._media = media or {}
        self._raise = export_raises or {}
        # file_list: list of file metadata dicts returned by files().list()
        self._file_list = file_list or []
        # Instrumentation for tests that need to assert de-dup: how many times
        # export()/get_media() was actually called per fileId.
        self.export_calls: dict[str, int] = {}
        self.get_media_calls: dict[str, int] = {}

    def export(self, fileId, mimeType):
        # Mirror the real Drive v3 API: files.export does NOT accept
        # supportsAllDrives. Passing it would raise TypeError here (as in prod),
        # so this signature guards against the kwarg being re-added.
        self.export_calls[fileId] = self.export_calls.get(fileId, 0) + 1
        if fileId in self._raise:
            return _Req(raise_exc=self._raise[fileId])
        return _Req(self._exports.get(fileId, b""))

    def get_media(self, fileId, supportsAllDrives=None):
        assert supportsAllDrives is True, (
            "get_media() must pass supportsAllDrives=True — required by the real "
            "Drive v3 API for files inside a Shared Drive"
        )
        self.get_media_calls[fileId] = self.get_media_calls.get(fileId, 0) + 1
        return _Req(self._media.get(fileId, b""))

    def list(self, **_kw):
        return _Req({"files": self._file_list})


class _Drives:
    def __init__(self, drives=None):
        self._drives = drives or []

    def list(self, **_kw):
        return _Req({"drives": self._drives})


class FakeDriveService:
    def __init__(self, **kw):
        # initial_cursor is the pageToken the first delta call will carry.
        # Defaults to start_token so the most common case (cursor=="100")
        # routes correctly without needing to pass it explicitly.
        start = kw.get("start_token", "100")
        initial = kw.get("initial_cursor", start)
        self._changes = _Changes(start, kw.get("pages"), initial_cursor=initial)
        self._files = _Files(
            kw.get("exports"),
            kw.get("media"),
            kw.get("export_raises"),
            kw.get("file_list"),
        )
        self._drives = _Drives(kw.get("shared_drives"))

    def changes(self):
        return self._changes

    def files(self):
        return self._files

    def drives(self):
        return self._drives


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _gdoc_change(fid, name="Doc", removed=False):
    ch = {"fileId": fid, "removed": removed}
    if not removed:
        ch["file"] = {
            "id": fid,
            "name": name,
            "mimeType": "application/vnd.google-apps.document",
            "modifiedTime": "2026-05-01T10:00:00Z",
            "owners": [{"displayName": "Someone"}],
        }
    return ch


def _plain_change(fid, name="Note", mime="text/plain"):
    return {
        "fileId": fid,
        "removed": False,
        "file": {
            "id": fid,
            "name": name,
            "mimeType": mime,
            "modifiedTime": "2026-05-01T10:00:00Z",
            "owners": [{"displayName": "Owner"}],
        },
    }


def _page(changes, next_page_token=None, new_start_page_token=None):
    p = {"changes": changes}
    if next_page_token is not None:
        p["nextPageToken"] = next_page_token
    if new_start_page_token is not None:
        p["newStartPageToken"] = new_start_page_token
    return p


def _store(tmp_path):
    s = Store(tmp_path / "test.sqlite3", dim=4)
    s.init()
    return s


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

def test_bootstrap_sets_cursor_no_files(tmp_path):
    """First run: no cursor. getStartPageToken returns "100".
    sync_drive returns 0, cursor is set to "100", no chunks upserted."""
    store = _store(tmp_path)
    svc = FakeDriveService(start_token="100")

    result = sync_drive(svc, store)

    assert result == 0
    assert store.get_cursor("drive") == "100"
    assert store.unembedded_chunks() == []


def test_delta_google_doc_exported_and_upserted(tmp_path):
    """Delta run: cursor "100", one Google Doc change, text exported and upserted.
    Cursor advances to "105", return value 1, chunk present."""
    store = _store(tmp_path)
    store.set_cursor("drive", "100")

    pages = [
        _page(
            [_gdoc_change("f1", "Budget Plan")],
            new_start_page_token="105",
        )
    ]
    svc = FakeDriveService(
        pages=pages,
        exports={"f1": b"Budget plan for Q3"},
    )

    result = sync_drive(svc, store)

    assert result == 1
    assert store.get_cursor("drive") == "105"
    chunk = store.get_chunk("gdrive-f1-0")
    assert chunk is not None
    assert "Budget plan" in chunk["text"]


def test_text_file_via_get_media(tmp_path):
    """text/plain file fetched via get_media, upserted as gdrive-f2-0."""
    store = _store(tmp_path)
    store.set_cursor("drive", "100")

    pages = [
        _page(
            [_plain_change("f2", "Meeting Notes", "text/plain")],
            new_start_page_token="106",
        )
    ]
    svc = FakeDriveService(
        pages=pages,
        media={"f2": b"meeting notes here"},
    )

    result = sync_drive(svc, store)

    assert result == 1
    chunk = store.get_chunk("gdrive-f2-0")
    assert chunk is not None
    assert "meeting notes" in chunk["text"]


def test_removed_change_skipped(tmp_path):
    """A change with removed=True is not upserted and not counted."""
    store = _store(tmp_path)
    store.set_cursor("drive", "100")

    removed_change = {"fileId": "f3", "removed": True}
    pages = [
        _page([removed_change], new_start_page_token="107")
    ]
    svc = FakeDriveService(pages=pages)

    result = sync_drive(svc, store)

    assert result == 0
    assert store.get_chunk("gdrive-f3-0") is None


def test_unsupported_mime_skipped(tmp_path):
    """image/png file: _fetch_text returns None, not upserted, not counted."""
    store = _store(tmp_path)
    store.set_cursor("drive", "100")

    img_change = {
        "fileId": "f4",
        "removed": False,
        "file": {
            "id": "f4",
            "name": "photo.png",
            "mimeType": "image/png",
            "modifiedTime": "2026-05-01T10:00:00Z",
            "owners": [],
        },
    }
    pages = [
        _page([img_change], new_start_page_token="108")
    ]
    svc = FakeDriveService(pages=pages)

    result = sync_drive(svc, store)

    assert result == 0
    assert store.get_chunk("gdrive-f4-0") is None


def test_pagination_processes_all(tmp_path):
    """Two pages: page 0 has nextPageToken -> page 1; page 1 has newStartPageToken.
    Files on both pages are upserted; cursor equals last newStartPageToken."""
    store = _store(tmp_path)
    store.set_cursor("drive", "0")  # '0' maps to pages[0]

    pages = [
        # page 0: index 0, nextPageToken "1" -> routes to pages[1]
        _page(
            [_gdoc_change("fa", "Doc A")],
            next_page_token="1",
        ),
        # page 1: index 1, last page carries newStartPageToken
        _page(
            [_gdoc_change("fb", "Doc B")],
            new_start_page_token="200",
        ),
    ]
    svc = FakeDriveService(
        pages=pages,
        exports={
            "fa": b"Content of Doc A for testing",
            "fb": b"Content of Doc B for testing",
        },
    )

    result = sync_drive(svc, store)

    assert result == 2
    assert store.get_chunk("gdrive-fa-0") is not None
    assert store.get_chunk("gdrive-fb-0") is not None
    assert store.get_cursor("drive") == "200"


def test_cursor_not_advanced_on_fetch_error(tmp_path):
    """If export raises RuntimeError, sync_drive propagates it and cursor stays unchanged."""
    store = _store(tmp_path)
    store.set_cursor("drive", "100")

    pages = [
        _page(
            [_gdoc_change("f5", "Failing Doc")],
            new_start_page_token="110",
        )
    ]
    svc = FakeDriveService(
        pages=pages,
        export_raises={"f5": RuntimeError("export failed")},
    )

    with pytest.raises(RuntimeError, match="export failed"):
        sync_drive(svc, store)

    assert store.get_cursor("drive") == "100"


# ---------------------------------------------------------------------------
# Unit tests for helpers
# ---------------------------------------------------------------------------

def test_fetch_text_google_doc():
    """_fetch_text routes Google Doc to export and returns decoded text."""
    meta = {"id": "x1", "mimeType": "application/vnd.google-apps.document"}
    svc = FakeDriveService(exports={"x1": b"Hello world"})
    assert _fetch_text(svc, meta) == "Hello world"


def test_fetch_text_plain_via_get_media():
    """_fetch_text routes text/plain to get_media."""
    meta = {"id": "x2", "mimeType": "text/plain"}
    svc = FakeDriveService(media={"x2": b"plain text content"})
    assert _fetch_text(svc, meta) == "plain text content"


def test_fetch_text_image_still_skipped():
    """_fetch_text returns None for image/png — images are not extracted."""
    meta = {"id": "x3", "mimeType": "image/png"}
    svc = FakeDriveService()
    assert _fetch_text(svc, meta) is None


def test_normalise_drive_produces_correct_doc_ids():
    """normalise_drive: doc_id pattern is gdrive-<id>-<i>; metadata has expected fields."""
    meta = {
        "id": "abc123",
        "name": "Test File",
        "mimeType": "application/vnd.google-apps.document",
        "modifiedTime": "2026-05-01T10:00:00Z",
        "owners": [{"displayName": "Test Owner"}],
    }
    chunks = normalise_drive(meta, "Some meaningful content for testing chunking behaviour here.")

    assert len(chunks) >= 1
    assert chunks[0].doc_id == "gdrive-abc123-0"
    assert chunks[0].metadata["source_type"] == "gdrive"
    assert chunks[0].metadata["file_id"] == "abc123"
    assert chunks[0].metadata["owner"] == "Test Owner"


def test_normalise_drive_empty_text_returns_empty():
    """normalise_drive returns [] for empty or whitespace-only text."""
    meta = {"id": "z1", "name": "Empty", "mimeType": "text/plain"}
    assert normalise_drive(meta, "") == []
    assert normalise_drive(meta, "   \n  ") == []


# ---------------------------------------------------------------------------
# Binary extractor integration tests
# ---------------------------------------------------------------------------

def _make_docx_bytes() -> bytes:
    import io
    from docx import Document
    doc = Document()
    doc.add_paragraph("Quarterly budget review")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Revenue"
    table.rows[0].cells[1].text = "Expenses"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_fetch_text_docx_via_get_media(tmp_path):
    """DOCX file: _fetch_text fetches via get_media and extracts text.
    Via backfill_drive it upserts a gdrive-<id>-0 chunk."""
    docx_bytes = _make_docx_bytes()
    DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    # _fetch_text unit check
    meta = {"id": "d1", "mimeType": DOCX_MIME}
    svc = FakeDriveService(media={"d1": docx_bytes})
    text = _fetch_text(svc, meta)
    assert text is not None
    assert "Quarterly budget review" in text
    assert "Revenue" in text

    # Integration: backfill_drive upserts the chunk
    store = _store(tmp_path)
    fmeta = {
        "id": "d1",
        "name": "Budget.docx",
        "mimeType": DOCX_MIME,
        "modifiedTime": "2026-05-01T10:00:00Z",
        "owners": [{"displayName": "Sam"}],
    }
    svc2 = FakeDriveService(
        media={"d1": docx_bytes},
        file_list=[fmeta],
    )
    processed = backfill_drive(svc2, store, "2026-01-01T00:00:00Z")
    assert processed == 1
    chunk = store.get_chunk("gdrive-d1-0")
    assert chunk is not None
    assert "Quarterly budget review" in chunk["text"]


def test_fetch_text_sheets_export_csv():
    """Google Sheets file: _fetch_text uses export(mimeType='text/csv'); text returned."""
    SHEETS_MIME = "application/vnd.google-apps.spreadsheet"
    csv_bytes = b"Month,Revenue\nJanuary,50000\nFebruary,62000\n"

    meta = {"id": "s1", "mimeType": SHEETS_MIME}
    svc = FakeDriveService(exports={"s1": csv_bytes})
    text = _fetch_text(svc, meta)
    assert text is not None
    assert "Month" in text
    assert "January" in text
    assert "50000" in text


# ---------------------------------------------------------------------------
# Task 2 duty-cycle fix: budget-interrupted mid-upsert must checkpoint safely
# ---------------------------------------------------------------------------

class _FakeBudget:
    """expired() returns False for the first `expire_after_calls` calls, True
    from then on — pins EXACTLY which iteration a real Budget's wall-clock
    expiry would have landed on, deterministically."""

    def __init__(self, expire_after_calls):
        self.calls = 0
        self.expire_after_calls = expire_after_calls

    def expired(self) -> bool:
        self.calls += 1
        return self.calls > self.expire_after_calls


def test_budget_interrupted_mid_upsert_resumes_without_skip_or_duplicate(tmp_path):
    """Mirrors test_gmail_sync.py's checkpoint-resume test for the Drive delta
    path. `newStartPageToken` is only emitted once ALL pages are consumed, so
    (like Gmail's historyId) advancing the cursor before every pending file
    is durably upserted would silently skip whatever wasn't reached yet.

    Verifies: (1) an interrupted run upserts only the files it reached and
    leaves the cursor at its OLD value; (2) a follow-up call with no budget
    completes the resume — f1 (already durably done in the first call) is
    genuinely SKIPPED in the upsert loop via the persisted `drive:resume_ids`
    set, only f2 is newly processed — and the cursor advances to the true
    newStartPageToken. (Unlike Gmail, the PAGINATION loop's own `_fetch_text`
    calls are not gated on the resume set here — see sync_drive's docstring
    for why that's an accepted, documented trade-off — so f1's export IS
    still re-fetched during pagination on resume; only the upsert/count is
    genuinely incremental.)
    """
    store = _store(tmp_path)
    store.set_cursor("drive", "100")

    pages = [
        _page(
            [_gdoc_change("f1", "Doc One"), _gdoc_change("f2", "Doc Two")],
            new_start_page_token="105",
        )
    ]
    svc = FakeDriveService(
        pages=pages,
        exports={"f1": b"first document content, long enough to matter",
                "f2": b"second document content, long enough to matter"},
    )

    # Call 1: pagination check before the (only) page (not expired -> both
    # f1/f2 fetched+buffered into `pending` during that one page). f1 is then
    # written WITHOUT consulting the budget -- the minimum-forward-progress
    # guarantee (see sync_drive): a budget already spent by the fetch phase must
    # still yield one write, or the resume set never grows and the round
    # livelocks. Call 2: upsert-loop check before f2 (expired -> stop).
    budget = _FakeBudget(expire_after_calls=1)
    result = sync_drive(svc, store, budget=budget)

    assert result == 1, "only the file(s) upserted before budget expiry should count"
    assert store.get_cursor("drive") == "100", (
        "cursor must NOT advance on a partial run — an early advance would "
        "silently and permanently skip f2 (newStartPageToken is only valid "
        "once every pending file from this delta window is durable)"
    )
    assert store.get_chunk("gdrive-f1-0") is not None
    assert store.get_chunk("gdrive-f2-0") is None

    # Resume: cursor unchanged, so pagination re-lists the SAME page (and
    # re-fetches f1's/f2's text — see docstring), but the upsert loop must
    # SKIP f1 (already resumed) and process only f2.
    result2 = sync_drive(svc, store, budget=None)

    assert result2 == 1, "only the genuinely new file (f2) should be counted/upserted on resume"
    assert store.get_cursor("drive") == "105"
    assert store.get_chunk("gdrive-f1-0") is not None
    assert store.get_chunk("gdrive-f2-0") is not None
    assert store.get_cursor("drive:resume_ids") == "[]", "resume set must be cleared once the round closes"

    # No duplicate-visible-effect regardless: f1 stayed one row (upsert keyed
    # on doc_id), never inserted twice.
    with store._connect() as db:
        count = db.execute(
            "SELECT COUNT(*) FROM chunks WHERE doc_id='gdrive-f1-0'"
        ).fetchone()[0]
    assert count == 1


def test_budget_interrupted_across_many_cycles_eventually_completes(tmp_path):
    """Critical-B reproduction, My-Drive variant (adversarial review, Task 2
    round 3): a delta bigger than one budget's worth of files must not
    livelock. Drives a 7-file delta through repeated budget-truncated calls
    (2 files' worth of upsert capacity each) and asserts the cursor
    eventually reaches the true final newStartPageToken and every file is
    durably upserted exactly once."""
    store = _store(tmp_path)
    store.set_cursor("drive", "100")

    n = 7
    changes = [_gdoc_change(f"f{i}", f"Doc {i}") for i in range(1, n + 1)]
    exports = {f"f{i}": f"document body number {i}, long enough to matter".encode()
              for i in range(1, n + 1)}
    pages = [_page(changes, new_start_page_token="200")]
    svc = FakeDriveService(pages=pages, exports=exports)

    per_call_capacity = 2
    max_cycles = 20
    for _cycle in range(max_cycles):
        if store.get_cursor("drive") != "100":
            break
        budget = _FakeBudget(expire_after_calls=1 + per_call_capacity)
        sync_drive(svc, store, budget=budget)
    else:
        raise AssertionError(
            f"cursor never advanced past the original delta window after "
            f"{max_cycles} cycles — this is the livelock the fix targets"
        )

    assert store.get_cursor("drive") == "200"
    assert store.get_cursor("drive:resume_ids") == "[]"
    for i in range(1, n + 1):
        doc_id = f"gdrive-f{i}-0"
        assert store.get_chunk(doc_id) is not None, f"f{i} was never ingested"
        with store._connect() as db:
            count = db.execute(
                "SELECT COUNT(*) FROM chunks WHERE doc_id=?", (doc_id,)
            ).fetchone()[0]
        assert count == 1, f"f{i} produced more than one chunk row"


def test_file_edited_mid_round_is_picked_up_not_skipped(tmp_path):
    """New Critical found in adversarial review round 4: once a file's id
    landed in the resume set (round 3's fix), it was skipped for the REST OF
    THAT ROUND no matter what -- including if the file changed in between.
    The round then closed and the real cursor advanced PAST the file's
    change record, with nothing left to re-surface it until the file
    happened to change again after the cursor had already moved on.

    Reproduced directly before this fix (against the round-3 code): an
    edited file's stored text stayed at its pre-edit content forever after
    the round closed. Fixed by keying the resume set on id+version
    (_file_resume_key: md5Checksum, or version+modifiedTime), not bare id,
    so an edit produces a DIFFERENT key and is recognized as new work rather
    than matched against the stale resume entry.
    """
    store = _store(tmp_path)
    store.set_cursor("drive", "100")

    f1 = _gdoc_change("f1", "Doc One")
    f1["file"]["md5Checksum"] = "hash-v1"
    f2 = _gdoc_change("f2", "Doc Two")
    f2["file"]["md5Checksum"] = "hash-v1-f2"
    pages = [_page([f1, f2], new_start_page_token="200")]
    svc = FakeDriveService(pages=pages, exports={
        "f1": b"ORIGINAL f1 body content, long enough to matter",
        "f2": b"f2 body content, long enough to matter",
    })

    upsert_calls = []
    orig_upsert = store.upsert_chunk

    def spy_upsert(*a, **kw):
        upsert_calls.append(a[0])  # doc_id
        return orig_upsert(*a, **kw)

    store.upsert_chunk = spy_upsert

    # Call 1: budget cuts off right after f1 is processed with its ORIGINAL
    # content -- f1's (stale-version) key lands in the resume set. f1 itself is
    # written unconditionally under the minimum-forward-progress guarantee, so
    # the cut-off lands one expired() call earlier than it used to.
    budget = _FakeBudget(expire_after_calls=1)
    sync_drive(svc, store, budget=budget)
    assert store.get_cursor("drive") == "100", "round must still be open"
    assert store.get_chunk("gdrive-f1-0")["text"].startswith("ORIGINAL")

    # f1 is edited in Drive (content AND version change) WHILE the round is
    # still open.
    svc._files._exports["f1"] = b"REVISED f1 body content, long enough to matter"
    f1_edited = _gdoc_change("f1", "Doc One")
    f1_edited["file"]["md5Checksum"] = "hash-v2"
    svc._changes._pages[0] = {"changes": [f1_edited, f2], "newStartPageToken": "200"}

    # Call 2: unbounded, completes the round.
    sync_drive(svc, store, budget=None)

    assert store.get_cursor("drive") == "200", "round must close"
    assert store.get_chunk("gdrive-f1-0")["text"].startswith("REVISED"), (
        "f1's edit must land -- the resume set must not have permanently "
        "skipped it just because its OLD id+version key was already "
        "resumed from call 1"
    )
    assert store.get_cursor("drive:resume_ids") == "[]"

    # f1 was upserted exactly twice total across the two calls (once with
    # its original content, once with the edit) -- proving the edit is
    # picked up exactly once per version, not repeatedly re-applied within
    # the same round nor silently dropped.
    assert upsert_calls.count("gdrive-f1-0") == 2


_EXPORTS_3 = {f"f{i}": (f"document {i} content, long enough to matter".encode())
              for i in range(3)}


def test_budget_spent_during_downloads_still_makes_forward_progress(tmp_path):
    """A budget spent inside the FETCH phase must not livelock sync_drive.

    _fetch_text runs for every changed file inside the pagination loop, with no
    budget check and no resume-set consultation. The write loop then checks the
    budget BEFORE its first item, so if the downloads spent it, zero items are
    written, resumed_ids never grows, and the cursor never advances -- the next
    cycle re-downloads exactly the same files, forever. Reproduced over 6
    consecutive cycles: processed=0, cursor frozen, 3 downloads re-issued each
    time.

    Two guarantees are asserted: at least one item is written per call (so the
    resume set always grows), and repeated cycles eventually complete the round.
    """
    store = _store(tmp_path)
    store.set_cursor("drive", "100")
    pages = [_page([_gdoc_change(f"f{i}", f"Doc {i}") for i in range(3)],
                   new_start_page_token="105")]

    total = 0
    for cycle in range(6):
        svc = FakeDriveService(pages=list(pages), start_token="100",
                               exports=_EXPORTS_3)
        # Survives the pagination check (call 1), then dies before the write
        # loop's first item -- i.e. consumed by the per-file downloads that
        # happen between the two. Expiring at call 0 instead would mean the
        # page is never even fetched, where doing nothing IS correct.
        got = sync_drive(svc, store, budget=_FakeBudget(expire_after_calls=1))
        total += got
        if store.get_cursor("drive") == "105":
            break

    assert total >= 1, (
        "no forward progress across 6 cycles with an immediately-expired "
        "budget -- this is the livelock"
    )
    assert store.get_cursor("drive") == "105", "round never closed"


def test_already_resumed_files_are_not_re_downloaded(tmp_path, monkeypatch):
    """Skipping must happen BEFORE _fetch_text, not after: otherwise a truncated
    cycle re-issues network exports for files it already durably wrote."""
    from mcpbrain.sync import drive as drive_mod
    store = _store(tmp_path)
    store.set_cursor("drive", "100")
    pages = [_page([_gdoc_change(f"f{i}", f"Doc {i}") for i in range(3)],
                   new_start_page_token="105")]

    calls = []
    real = drive_mod._fetch_text
    monkeypatch.setattr(drive_mod, "_fetch_text",
                        lambda svc, fmeta: calls.append(fmeta.get("id")) or real(svc, fmeta))

    # First cycle: budget allows exactly one write, so 2 files stay pending.
    svc = FakeDriveService(pages=list(pages), start_token="100",
                           exports=_EXPORTS_3)
    sync_drive(svc, store, budget=_FakeBudget(expire_after_calls=1))
    first_round = len(calls)
    assert first_round == 3, "all three should be fetched on the first pass"

    calls.clear()
    svc2 = FakeDriveService(pages=list(pages), start_token="100",
                            exports=_EXPORTS_3)
    sync_drive(svc2, store, budget=_FakeBudget(expire_after_calls=1))
    assert len(calls) < first_round, (
        f"re-downloaded {len(calls)} of {first_round} already-checkpointed files"
    )


# ---------------------------------------------------------------------------
# Gate 3 / Task 4: fetch_content, folder_path, upsert_file_chunks
# ---------------------------------------------------------------------------

def test_an_unsupported_drive_type_is_recorded_rather_than_silently_dropped():
    """A2: .pptx, .doc, .pages, images and .zip all returned None from
    _fetch_text with no chunk, no stub and no log line."""
    from mcpbrain.sync import drive

    class _Store:
        def __init__(self):
            self.changes = []

        def record_change(self, kind, ref_id="", summary=""):
            self.changes.append((kind, ref_id, summary))

    store = _Store()
    fmeta = {"id": "f-1", "name": "Deck.key",
             "mimeType": "application/x-iwork-keynote-sffkey"}

    assert drive.fetch_content(object(), fmeta, store=store) is None
    assert store.changes and store.changes[0][0] == "ingest_skip"
    assert "unsupported_mime" in store.changes[0][2]


def test_a_supported_type_that_extracts_to_nothing_is_recorded_distinctly(monkeypatch):
    """B7: eight `except Exception: return ""` sites make a corrupt DOCX
    indistinguishable from an unsupported type. They must not share a bucket."""
    from mcpbrain.sync import drive

    class _Store:
        def __init__(self):
            self.changes = []

        def record_change(self, kind, ref_id="", summary=""):
            self.changes.append((kind, ref_id, summary))

    store = _Store()
    monkeypatch.setattr(drive, "_fetch_text", lambda service, meta: "")
    fmeta = {"id": "f-2", "name": "Broken.docx",
             "mimeType": "application/vnd.openxmlformats-officedocument."
                         "wordprocessingml.document"}

    drive.fetch_content(object(), fmeta, store=store)

    assert [s.split(":")[0] for _k, _r, s in store.changes] == ["extraction_empty"]


def test_folder_path_is_resolved_and_cached():
    """C5: embed.contextual_prefix reads metadata['folder_path'] and
    normalise_drive never wrote it, so every Drive contextual prefix has been
    missing its folder context — dead provenance in a default-ON feature."""
    from mcpbrain.sync import drive

    calls: list = []

    class _Service:
        def files(self):
            return self

        def get(self, fileId, fields, supportsAllDrives=None):
            calls.append(fileId)
            self._fid = fileId
            return self

        def execute(self):
            return {"folder-1": {"id": "folder-1", "name": "Budgets",
                                 "parents": ["folder-0"]},
                    "folder-0": {"id": "folder-0", "name": "Finance",
                                 "parents": []}}[self._fid]

    cache: dict = {}
    fmeta = {"id": "f1", "name": "Budget.xlsx", "parents": ["folder-1"]}

    assert drive.folder_path(_Service(), fmeta, cache) == "Finance/Budgets"
    drive.folder_path(_Service(), fmeta, cache)
    assert calls == ["folder-1", "folder-0"], "the second call must hit the cache"


def test_a_shrinking_document_drops_its_orphaned_chunks(tmp_path):
    """B5: Drive writes gdrive-<fid>-<i> for i in 0..n-1 and only ever upserts.
    Nothing deleted indices n..m left by a previous, longer version, so deleted
    paragraphs stayed searchable indefinitely and were re-fed to expansion as
    current content."""
    from mcpbrain.store import Store
    from mcpbrain.sync.drive import normalise_drive, upsert_file_chunks

    store = Store(tmp_path / "b.sqlite3", dim=4)
    store.init()
    fmeta = {"id": "f1", "name": "Notes.txt", "mimeType": "text/plain"}

    long_text = "\n\n".join(f"Para {i} " + "word " * 400 for i in range(5))
    upsert_file_chunks(store, normalise_drive(fmeta, long_text), file_id="f1")
    assert len(store.doc_ids_for_file("f1")) >= 3

    upsert_file_chunks(store, normalise_drive(fmeta, "Para 0 " + "word " * 100),
                       file_id="f1")

    assert store.doc_ids_for_file("f1") == ["gdrive-f1-0"], "stale chunks survived"


def test_upserting_an_unchanged_document_deletes_nothing(tmp_path):
    from mcpbrain.store import Store
    from mcpbrain.sync.drive import normalise_drive, upsert_file_chunks

    store = Store(tmp_path / "b.sqlite3", dim=4)
    store.init()
    fmeta = {"id": "f1", "name": "Notes.txt", "mimeType": "text/plain"}
    text = "\n\n".join(f"Para {i} " + "word " * 400 for i in range(5))

    upsert_file_chunks(store, normalise_drive(fmeta, text), file_id="f1")
    first = sorted(store.doc_ids_for_file("f1"))
    upsert_file_chunks(store, normalise_drive(fmeta, text), file_id="f1")

    assert sorted(store.doc_ids_for_file("f1")) == first


# ---------------------------------------------------------------------------
# Post-approval review finding: fetch_content's per-file record_skip floods
# change_log. Aggregated over a sync round via a `report` dict instead.
# ---------------------------------------------------------------------------

def test_fetch_content_report_tallies_instead_of_writing_immediately():
    """Unit-level: passing `report=` must switch fetch_content from an
    immediate store.record_change per call to tallying {(kind, mime): count}
    in the caller-owned dict, with nothing written to the store at all."""
    from mcpbrain.sync import drive

    class _Store:
        def __init__(self):
            self.changes = []

        def record_change(self, kind, ref_id="", summary=""):
            self.changes.append((kind, ref_id, summary))

    store = _Store()
    report: dict = {}
    fmeta_png = {"id": "f-1", "name": "a.png", "mimeType": "image/png"}
    fmeta_jpg = {"id": "f-2", "name": "b.jpg", "mimeType": "image/jpeg"}

    for _ in range(3):
        drive.fetch_content(object(), fmeta_png, store=store, report=report)
    drive.fetch_content(object(), fmeta_jpg, store=store, report=report)

    assert store.changes == [], "report= must suppress the immediate write"
    assert report == {("unsupported_mime", "image/png"): 3,
                      ("unsupported_mime", "image/jpeg"): 1}


def test_flush_skip_report_emits_one_bounded_row_per_kind():
    """flush_skip_report must turn a multi-mime tally into one change_log row
    per `kind`, with the per-mime breakdown folded into the detail text —
    not one row per (kind, mime) and definitely not one row per file."""
    from mcpbrain.sync import drive

    class _Store:
        def __init__(self):
            self.changes = []

        def record_change(self, kind, ref_id="", summary=""):
            self.changes.append((kind, ref_id, summary))

    store = _Store()
    report = {("unsupported_mime", "image/png"): 270,
              ("unsupported_mime", "image/jpeg"): 70,
              ("extraction_empty", "application/pdf"): 2}

    drive.flush_skip_report(store, report)

    assert len(store.changes) == 2, "one row per kind, not per (kind, mime)"
    by_kind = {c[2].split(":")[0]: c[2] for c in store.changes}
    assert "drive_unsupported_mime" in by_kind
    assert "270" in by_kind["drive_unsupported_mime"]
    assert "70" in by_kind["drive_unsupported_mime"]
    assert "image/png" in by_kind["drive_unsupported_mime"]
    assert "drive_extraction_empty" in by_kind
    assert "2" in by_kind["drive_extraction_empty"]


def test_flush_skip_report_is_a_noop_on_an_empty_report():
    from mcpbrain.sync import drive

    class _Store:
        def __init__(self):
            self.changes = []

        def record_change(self, kind, ref_id="", summary=""):
            self.changes.append((kind, ref_id, summary))

    store = _Store()
    drive.flush_skip_report(store, {})
    assert store.changes == []


def test_many_skipped_files_in_one_sync_round_produce_one_summary_row_not_one_per_file(tmp_path):
    """Review finding (post-Task-4-approval): fetch_content used to call
    ingest_report.record_skip once per skipped file inside sync_drive's
    listing loop — one store.record_change write, one change_log row, per
    file. change_log is pruned to 500 rows and doubles as the user-facing
    change digest (dashboard.py's recent_changes); a Drive sync whose window
    contains a few hundred images (a common real case) would evict the
    entire genuine audit trail and fill the digest with per-file noise.
    sync_drive must flush ONE aggregated row for the whole round instead."""
    store = _store(tmp_path)
    store.set_cursor("drive", "100")

    n = 50
    changes = [
        {"fileId": f"img{i}", "removed": False,
         "file": {"id": f"img{i}", "name": f"photo{i}.png", "mimeType": "image/png",
                  "modifiedTime": "2026-05-01T10:00:00Z", "owners": []}}
        for i in range(n)
    ]
    pages = [_page(changes, new_start_page_token="900")]
    svc = FakeDriveService(pages=pages)

    result = sync_drive(svc, store)

    assert result == 0, "every file is an unsupported image; none should count as processed"
    skip_rows = [c for c in store.recent_changes(limit=1000) if c["change_type"] == "ingest_skip"]
    assert len(skip_rows) == 1, (
        f"expected exactly one aggregated ingest_skip row for {n} skipped "
        f"files, got {len(skip_rows)} — the per-file flood is back"
    )
    assert "unsupported_mime" in skip_rows[0]["summary"]
    assert "image/png" in skip_rows[0]["summary"]
    assert str(n) in skip_rows[0]["summary"]


def test_backfill_drive_also_aggregates_skips_across_its_bounded_window(tmp_path):
    """Same finding, backfill_drive path: it advances no cursor and re-lists
    the same window every call, so per-file recording would re-flood on every
    single re-run. Confirms the aggregation applies there too."""
    store = _store(tmp_path)

    n = 12
    file_list = [
        {"id": f"img{i}", "name": f"photo{i}.png", "mimeType": "image/png",
         "modifiedTime": "2026-05-01T10:00:00Z", "owners": []}
        for i in range(n)
    ]
    svc = FakeDriveService(file_list=file_list)

    processed = backfill_drive(svc, store, "2020-01-01T00:00:00Z")

    assert processed == 0
    skip_rows = [c for c in store.recent_changes(limit=1000) if c["change_type"] == "ingest_skip"]
    assert len(skip_rows) == 1, f"expected one aggregated row, got {len(skip_rows)}"
    assert str(n) in skip_rows[0]["summary"]


# ---------------------------------------------------------------------------
# I9: a PARTIAL extraction must not trigger the B5 orphan-delete sweep.
# ---------------------------------------------------------------------------

def test_a_partial_extraction_does_not_delete_the_chunks_it_never_reached(tmp_path):
    """I9: extract_tables_from_xlsx / _xls / extract_text_from_pptx keep whatever
    they had when an exception hits mid-iteration (better than nothing). But
    upsert_file_chunks read the SHORT chunk list as evidence that the document had
    SHRUNK and deleted the higher-index "orphans" — so a transient failure on
    sheet 3 of 5 permanently deleted sheets 3-5's previously-good chunks, and
    nothing re-triggers extraction for a file whose metadata never changes again.
    A logged warning became irreversible content loss."""
    from mcpbrain.store import Store
    from mcpbrain.sync.drive import normalise_drive, upsert_file_chunks

    store = Store(tmp_path / "b.sqlite3", dim=4)
    store.init()
    fmeta = {"id": "f1", "name": "Budget.xlsx", "mimeType": "text/plain"}

    full = "\n\n".join(f"Sheet {i} " + "word " * 400 for i in range(5))
    upsert_file_chunks(store, normalise_drive(fmeta, full), file_id="f1")
    before = sorted(store.doc_ids_for_file("f1"))
    assert len(before) >= 3

    # The next round's extraction dies after sheet 1: a much shorter document.
    partial_chunks = normalise_drive(fmeta, "Sheet 0 " + "word " * 100)
    deleted = upsert_file_chunks(store, partial_chunks, file_id="f1", partial=True)

    assert deleted == 0
    assert sorted(store.doc_ids_for_file("f1")) == before, (
        "a partial extraction deleted the chunks it never reached"
    )


def test_a_complete_short_extraction_still_sweeps_orphans(tmp_path):
    """The discriminator: partial=False keeps B5's behaviour exactly."""
    from mcpbrain.store import Store
    from mcpbrain.sync.drive import normalise_drive, upsert_file_chunks

    store = Store(tmp_path / "b.sqlite3", dim=4)
    store.init()
    fmeta = {"id": "f1", "name": "Notes.txt", "mimeType": "text/plain"}

    full = "\n\n".join(f"Para {i} " + "word " * 400 for i in range(5))
    upsert_file_chunks(store, normalise_drive(fmeta, full), file_id="f1")

    upsert_file_chunks(store, normalise_drive(fmeta, "Para 0 " + "word " * 100),
                       file_id="f1", partial=False)

    assert store.doc_ids_for_file("f1") == ["gdrive-f1-0"]


def test_fetch_content_marks_a_partial_table_extraction(monkeypatch):
    """The signal has to survive the trip from the extractor to Content.partial,
    or the call sites can't act on it."""
    from mcpbrain.sync import drive
    from mcpbrain.sync.extractors import PartialTables
    from mcpbrain.sync.tabular import Table

    xlsx = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    class _Media:
        def execute(self):
            return b"fake"

    class _Files:
        def get_media(self, **kw):
            return _Media()

    class _Svc:
        def files(self):
            return _Files()

    monkeypatch.setattr(
        drive, "extract_tables_from_xlsx",
        lambda data, char_budget: PartialTables(
            [Table(sheet="S1", header=["a"], rows=[["1"]], rows_total=1)]))

    content = drive.fetch_content(_Svc(), {"id": "f1", "name": "b.xlsx",
                                           "mimeType": xlsx})

    assert content is not None and content.partial is True
    assert len(content.tables) == 1


def test_fetch_content_does_not_mark_a_complete_extraction(monkeypatch):
    from mcpbrain.sync import drive
    from mcpbrain.sync.tabular import Table

    xlsx = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    class _Media:
        def execute(self):
            return b"fake"

    class _Files:
        def get_media(self, **kw):
            return _Media()

    class _Svc:
        def files(self):
            return _Files()

    monkeypatch.setattr(
        drive, "extract_tables_from_xlsx",
        lambda data, char_budget: [Table(sheet="S1", header=["a"], rows=[["1"]],
                                         rows_total=1)])

    content = drive.fetch_content(_Svc(), {"id": "f1", "name": "b.xlsx",
                                           "mimeType": xlsx})

    assert content is not None and content.partial is False


def _partial_publish_harness(tmp_path, monkeypatch, *, partial: bool):
    """_cache_first_extract_one over a forced cache MISS, so the local-extraction
    path runs and its (file_id, content_hash) publish tuple is observable."""
    from mcpbrain import ingest_cache
    from mcpbrain.store import Store
    from mcpbrain.sync import drive

    store = Store(tmp_path / "b.sqlite3", dim=4)
    store.init()
    monkeypatch.setattr(ingest_cache, "try_import", lambda *a, **kw: False)
    monkeypatch.setattr(
        drive, "fetch_content",
        lambda *a, **kw: drive.Content(text="Para 0 " + "word " * 200,
                                       partial=partial))
    monkeypatch.setattr(drive, "folder_path", lambda *a, **kw: "")

    fmeta = {"id": "f1", "name": "Budget.xlsx", "mimeType": "text/plain",
             "md5Checksum": "abc123"}
    processed, miss = drive._cache_first_extract_one(
        object(), store, object(), "drv1", fmeta, {})
    return store, processed, miss


def test_a_partial_extraction_is_not_published_to_the_ingest_cache(tmp_path, monkeypatch):
    """The miss tuple _cache_first_extract_one returns is what the caller
    publishes as the fleet-wide ingest-cache artifact for that content hash. I9
    stopped a partial extraction from deleting chunks locally, but still published
    it — so a truncated document propagated to every other install under a hash
    that says it is complete, and would not self-heal until the file changed."""
    store, processed, miss = _partial_publish_harness(
        tmp_path, monkeypatch, partial=True)

    assert miss is None, "a truncated extraction was published fleet-wide"
    assert processed is True, (
        "the file WAS indexed locally — only the cache publish is suppressed")
    assert store.doc_ids_for_file("f1"), "local chunks must still be written"


def test_a_complete_extraction_is_still_published_to_the_ingest_cache(tmp_path, monkeypatch):
    """The discriminator: the normal path must keep publishing, or every install
    re-extracts every shared-drive file forever."""
    _store, processed, miss = _partial_publish_harness(
        tmp_path, monkeypatch, partial=False)

    assert processed is True
    assert miss is not None and miss[0] == "f1"


def test_the_aggregated_skip_row_names_the_source_it_came_from():
    """Please-fix minor: flush_skip_report passed ref_id="", so the aggregated
    rows could not be traced to the drive that produced them. Mirrors
    sync/gmail.py's reviewed pattern of passing `source` as ref_id."""
    from mcpbrain.sync import drive

    class _Store:
        def __init__(self):
            self.changes = []

        def record_change(self, kind, ref_id="", summary=""):
            self.changes.append((kind, ref_id, summary))

    store = _Store()
    drive.flush_skip_report(store, {("unsupported_mime", "image/png"): 3},
                            source="drive:DRIVE123")

    assert store.changes[0][1] == "drive:DRIVE123"


def test_reingest_files_replaces_a_files_chunks_from_a_fresh_fetch(tmp_path, monkeypatch):
    """There is no targeted re-ingest path: backfill_drive filters on
    modifiedTime, and the delta sync only sees CHANGED files — so a file whose
    content is fine but whose CHUNKING is out of date can never be revisited.
    455 clipped spreadsheets and 9,351 legacy files need exactly that."""
    from mcpbrain.store import Store
    from mcpbrain.sync import drive

    store = Store(tmp_path / "b.sqlite3", dim=4)
    store.init()
    # Pre-existing chunks from the old chunker: more of them, and stale text.
    for i in range(4):
        store.upsert_chunk(f"gdrive-f1-{i}", f"|  |  | old {i} |", f"h{i}",
                           {"source_type": "gdrive", "file_id": "f1",
                            "chunk_index": i})

    class _Service:
        def files(self):
            return self

        def get(self, fileId, fields=None, supportsAllDrives=None):
            self._fid = fileId
            return self

        def get_media(self, fileId, supportsAllDrives=None):
            return self

        def execute(self):
            return {"id": "f1", "name": "Notes.txt", "mimeType": "text/plain",
                    "modifiedTime": "2026-07-01T00:00:00Z", "parents": []}

    monkeypatch.setattr(drive, "_fetch_text",
                        lambda service, meta: "Recovered prose content.")

    summary = drive.reingest_files(_Service(), store, ["f1"])

    assert summary["files"] == 1
    remaining = sorted(store.doc_ids_for_file("f1"))
    assert remaining == ["gdrive-f1-0"], f"stale chunks survived: {remaining}"
    assert "Recovered" in store.get_chunk("gdrive-f1-0")["text"]


def test_reingest_files_skips_a_file_that_no_longer_exists(tmp_path):
    """A file deleted from Drive since it was chunked must not abort the run or
    delete its chunks — that is the removal path's job, not the repair's."""
    from googleapiclient.errors import HttpError

    from mcpbrain.store import Store
    from mcpbrain.sync import drive

    store = Store(tmp_path / "b.sqlite3", dim=4)
    store.init()
    store.upsert_chunk("gdrive-gone-0", "text", "h", {"source_type": "gdrive",
                                                      "file_id": "gone"})

    class _Resp:
        status = 404
        reason = "Not Found"

    class _Service:
        def files(self):
            return self

        def get(self, **kw):
            return self

        def execute(self):
            raise HttpError(_Resp(), b"not found")

    summary = drive.reingest_files(_Service(), store, ["gone"])

    assert summary["missing"] == 1
    assert summary["files"] == 0
    assert store.get_chunk("gdrive-gone-0") is not None


def test_reingest_files_isolates_a_non_404_httperror_from_metadata_fetch(tmp_path, monkeypatch):
    """Only a 404 (the file is genuinely gone) is special-cased as `missing`.
    A 403/429/5xx from files().get() itself -- all realistic across a
    9,351-file batch -- is a per-file failure like any other and must not
    escape the loop and abort the whole run; the next file must still be
    reached and processed."""
    from googleapiclient.errors import HttpError

    from mcpbrain.store import Store
    from mcpbrain.sync import drive

    store = Store(tmp_path / "b.sqlite3", dim=4)
    store.init()

    class _Resp:
        status = 429
        reason = "Too Many Requests"

    class _Service:
        def files(self):
            return self

        def get(self, fileId=None, **kw):
            self._fid = fileId
            return self

        def execute(self):
            if self._fid == "throttled":
                raise HttpError(_Resp(), b"rate limited")
            return {"id": self._fid, "name": f"{self._fid}.txt",
                    "mimeType": "text/plain", "parents": []}

    monkeypatch.setattr(drive, "_fetch_text", lambda service, meta: "fine content")

    summary = drive.reingest_files(_Service(), store, ["ok1", "throttled", "ok2"])

    assert summary["files"] == 2, "a non-404 HttpError must not abort the run"
    assert summary["failed"] == 1
    assert summary["missing"] == 0


def test_reingest_files_is_bounded_and_reports_per_file_failures(tmp_path, monkeypatch):
    """One unreadable file in 9,351 must not end the run."""
    from mcpbrain.store import Store
    from mcpbrain.sync import drive

    store = Store(tmp_path / "b.sqlite3", dim=4)
    store.init()

    class _Service:
        def files(self):
            return self

        def get(self, fileId=None, **kw):
            self._fid = fileId
            return self

        def execute(self):
            return {"id": self._fid, "name": f"{self._fid}.txt",
                    "mimeType": "text/plain", "parents": []}

    def _boom(service, meta):
        if meta["id"] == "bad":
            raise RuntimeError("extraction exploded")
        return "fine content"

    monkeypatch.setattr(drive, "_fetch_text", _boom)

    summary = drive.reingest_files(_Service(), store, ["ok1", "bad", "ok2"])

    assert summary["files"] == 2
    assert summary["failed"] == 1
